import streamlit as st
import docx
from docx import Document

menu = st.sidebar.header('Selecione o Programa abaixo:')
tela_principal = st.title(":gray[Olá bem seja bem vindo]")
st.markdown('Esta é uma aplicativo para fazer anotação das partes que serão apresentadas nas Assembléias e Congressos.')
st.markdown(':blue[No menu ao lado esquerdo selecione o programa de Assembléia ou Congresso para começar a fazer suas anotações]')

def asc():    
    if tela1:
        titulo = st.title(":blue[Programa da Assembleia de Circuito com o Superintendente de Circuito 2023 a 2024]")
        tema = st.header(":red[Espere Ansiosamente por Jeová!]\n Salmo 130:6")
        manha = st.header(':blue[Manhã]')

        Música = st.text("9:30 Música")
        cantico1 = st.text("9:40 Cântico 88 e oração")
        discurso = st.text_area("9:50 ‘Esperar ansiosamente por Jeová’ — Como?")
        serie_discurso1 = st.text_area("10:05 Série de discursos: \nImite aqueles que esperaram ansiosamente - Habacuque")
        serie_discurso2 = st.text_area("Imite aqueles que esperaram ansiosamente - João")
        serie_discurso3 = st.text_area("Imite aqueles que esperaram ansiosamente - Ana")
        cantico2 = st.text_input("11:05 Cântico 142 e anúncios")
        discurso4 = st.text_area("11:15 O que você está esperando?")
        discurso_batismo = st.text_area("11:30 Dedicação e batismo")
        cantico3 = st.text("12:00 Cântico 28")

        tarde = st.header(':blue[Tarde]')

        musica2 = st.text("13:10 Música")
        cantico4 = st.text("13:20 Cântico 54 e oração")
        discurso_publico = st.text_area("13:30 Discurso para o público: Será que a paciência ainda tem valor?")
        sentinela = st.text_area("14:00 Resumo de A Sentinela")
        cantico5 = st.text_input("14:30 Cântico 143 e anúncios")
        serie2_discurso1 = st.text_area("14:40 Série de discursos: Espere por Jeová \n Quando se sentir sozinho")
        serie2_discurso2 = st.text_area("Quando você cometer erros")
        serie2_discurso3 = st.text_area("Quando os maus forem bem-sucedidos")
        discurso5 = st.text_area("15:40 “Há uma recompensa para o justo")
        cantico6 = st.text("16:15 Cântico 140 e oração")

        st.title(':red[Preste atenção às respostas para as seguintes perguntas:]')

        recap_perg1 = st.text_area("1. Como nós ‘esperamos ansiosamente por Jeová’? (Salmo 130:5, 6)")
        recap_perg2 = st.text_area('2. Como podemos esperar ansiosamente por Jeová quando enfrentamos dificuldades? (Habacuque 2:3, 4; 2 Timóteo 4:2; Lucas 2:36-38')
        recap_perg3 = st.text_area('3. Como podemos usar de forma sábia o nosso tempo enquanto esperamos pelo dia de Jeová? (2 Pedro 3:11-13)')
        recap_perg4 = st.text_area('4. Por que podemos esperar por Jeová com confiança quando enfrentamos desafios? (Salmo 62:1, 2, 8, 10; Capítulo 68:6; Capítulo 130:2 a 4)')
        recap_perg5 = st.text_area('5. O que devemos fazer para receber a “recompensa para o justo”? (Salmo 58:11)')
        nome_arq = st.text_input('Nome do arquivo - no final do nome colocar (.doc)')

        if st.button('Salvar'):
            doc = Document()
            
            doc.add_heading('Programa da Assembleia de Circuito com o Superintendente de Circuito 2023 a 2024', 0)
            doc.add_heading('Espere Ansiosamente por Jeová! - Salmo 130:6', 1)
            doc.add_heading('Manhã', 0)
            doc.add_paragraph('9:30 Música')
            doc.add_paragraph(f'9:50 ‘Esperar ansiosamente por Jeová’ — Como? \n{discurso}')
            doc.add_paragraph(f'10:05 Série de discursos: \nImite aqueles que esperaram ansiosamente - Habacuque \n{serie_discurso1}')
            doc.add_paragraph(f'Imite aqueles que esperaram ansiosamente - João \n{serie_discurso2}')
            doc.add_paragraph(f'Imite aqueles que esperaram ansiosamente - Ana \n{serie_discurso3}')
            doc.add_paragraph(f'11:05 Cântico 142 e anúncios \n{cantico2}')
            doc.add_paragraph(f'11:15 O que você está esperando? \n{discurso4}')
            doc.add_paragraph(f'11:30 Dedicação e batismo \n{discurso_batismo}')
            doc.add_paragraph(f'12:00 Cântico 28')
            
            doc.add_heading('Tarde', 0)
            
            doc.add_paragraph(f'13:10 Música')
            doc.add_paragraph(f'13:20 Cântico 54 e oração')
            doc.add_paragraph(f'13:30 Discurso para o público: Será que a paciência ainda tem valor? \n{discurso_publico}')
            doc.add_paragraph(f'14:00 Resumo de A Sentinela \n{sentinela}')
            doc.add_paragraph(f'14:30 Cântico 143 e anúncios \n{cantico5}')
            doc.add_paragraph(f'14:40 Série de discursos: Espere por Jeová: \n Quando se sentir sozinho \n{serie2_discurso1}')
            doc.add_paragraph(f'Quando você cometer erros \n{serie2_discurso2}')
            doc.add_paragraph(f'Quando os maus forem bem-sucedidos \n{serie2_discurso3}')
            doc.add_paragraph(f'15:40 “Há uma recompensa para o justo \n{discurso5}')
            doc.add_paragraph(f'16:15 Cântico 140 e oração')
            doc.add_heading('Preste atenção às respostas para as seguintes perguntas:', 0)
            doc.add_paragraph(f'1. Como nós ‘esperamos ansiosamente por Jeová’? (Salmo 130:5, 6) \n{recap_perg1}')
            doc.add_paragraph(f'2. Como podemos esperar ansiosamente por Jeová quando enfrentamos dificuldades? (Habacuque 2:3, 4; 2 Timóteo 4:2; Lucas 2:36-38 \n{recap_perg2}')
            doc.add_paragraph(f'3. Como podemos usar de forma sábia o nosso tempo enquanto esperamos pelo dia de Jeová? (2 Pedro 3:11-13) \n{recap_perg3}')
            doc.add_paragraph(f'4. Por que podemos esperar por Jeová com confiança quando enfrentamos desafios? (Salmo 62:1, 2, 8, 10; Capítulo 68:6; Capítulo 130:2 a 4) \n{recap_perg4}')
            doc.add_paragraph(f'5. O que devemos fazer para receber a “recompensa para o justo”? (Salmo 58:11) \n{recap_perg5}')
        
        # Salva o documento Word
            doc_file = (nome_arq)
            doc.save(f'Docs\{nome_arq}')

            st.success(f"Documento Word gerado com sucesso, {doc_file}")


def arb():
    
    if tela2:
        titulo = st.title(":blue[Programa da Assembleia de Circuito de 2023-2024 com o Representante de Betel]")
        tema = st.header(":red[Entre no Descanso de Deus! — Hebreus 4:11]")
        manha = st.header(':blue[Manhã]')

        Música = st.text("9:40 Música")
        cantico1 = st.text("9:50 Cântico 87 e oração")
        discurso = st.text_area("10:15 “A palavra de Deus é viva” — O que isso significa?")
        serie_discurso1 = st.text_area("10:30 Continue a buscar as orientações de Jeová")
        # serie_discurso2 = st.text_area("Imite aqueles que esperaram ansiosamente - João")
        # serie_discurso3 = st.text_area("Imite aqueles que esperaram ansiosamente - Ana")
        cantico2 = st.text_input("10:55 Cântico 89 e anúncios")
        discurso4 = st.text_area("11:05 Jeová abençoa os obedientes")
        discurso_batismo = st.text_area("11:35 Dedicação e batismo")
        cantico3 = st.text("12:05 Cântico 32")

        tarde = st.header(':blue[Tarde]')

        musica2 = st.text("13:20 Música")
        cantico4 = st.text("13:30 Cântico 49")
        discurso_publico = st.text_area("13:35 Experiências")
        sentinela = st.text_area("13:45 Resumo de A Sentinela")
        serie2_discurso1 = st.text_area("14:15 Série de discursos: Eles alegram o coração de Jeová! \nJovens")
        serie2_discurso2 = st.text_area("Irmãs")
        serie2_discurso3 = st.text_area("Idosos")
        cantico5 = st.text_input("15:00 Cântico 38 e anúncios")
        discurso5 = st.text_area("15:10 Sinta alegria no seu serviço a Jeová")
        cantico6 = st.text("15:55 Cântico 118 e oração")

        st.title(':red[Preste atenção às respostas para as seguintes perguntas:]')

        recap_perg1 = st.text_area("1. Como entramos no descanso de Deus? (Gênesis 2:1-3; Hebreus 4:1, 11)")
        recap_perg2 = st.text_area('2. Como “a palavra de Deus” pode exercer poder em nossa vida? (1 Tessalonicenses 2:13; Hebreus 4:12)')
        recap_perg3 = st.text_area('3. Por que precisamos buscar as orientações de Jeová? (Isaías 26:7-9, 15, 20)')
        recap_perg4 = st.text_area('4. O que precisamos fazer para ser abençoados por Jeová? (1 Pedro 1:13-15; 1 João 5:3)')
        recap_perg5 = st.text_area('5. Como podemos alegrar o coração de Jeová? (Salmo 71:14, 15; Romanos 12:2; 1 Pedro 4:10)')
        recap_perg6 = st.text_area('6. Como podemos sentir alegria no nosso serviço a Jeová? (João 5:17)')
        nome_arq = st.text_input('Nome do arquivo - no final do nome colocar (.doc)')

        if st.button('Salvar'):
            doc = Document()
            
            doc.add_heading('Programa da Assembleia de Circuito com o Superintendente de Circuito 2023 a 2024', 0)
            doc.add_heading('Espere Ansiosamente por Jeová! - Salmo 130:6', 1)
            doc.add_heading('Manhã', 0)
            doc.add_paragraph('9:30 Música')
            doc.add_paragraph(f'9:50 ‘Esperar ansiosamente por Jeová’ — Como? \n{discurso}')
            doc.add_paragraph(f'10:05 Série de discursos: \nImite aqueles que esperaram ansiosamente - Habacuque \n{serie_discurso1}')
            # doc.add_paragraph(f'Imite aqueles que esperaram ansiosamente - João \n{serie_discurso2}')
            # doc.add_paragraph(f'Imite aqueles que esperaram ansiosamente - Ana \n{serie_discurso3}')
            doc.add_paragraph(f'11:05 Cântico 142 e anúncios \n{cantico2}')
            doc.add_paragraph(f'11:15 O que você está esperando? \n{discurso4}')
            doc.add_paragraph(f'11:30 Dedicação e batismo \n{discurso_batismo}')
            doc.add_paragraph(f'12:00 Cântico 28')
            
            doc.add_heading('Tarde', 0)
            
            doc.add_paragraph(f'13:10 Música')
            doc.add_paragraph(f'13:20 Cântico 54 e oração')
            doc.add_paragraph(f'13:30 Discurso para o público: Será que a paciência ainda tem valor? \n{discurso_publico}')
            doc.add_paragraph(f'14:00 Resumo de A Sentinela \n{sentinela}')
            doc.add_paragraph(f'14:30 Cântico 143 e anúncios \n{cantico5}')
            doc.add_paragraph(f'14:40 Série de discursos: Espere por Jeová: \n Quando se sentir sozinho \n{serie2_discurso1}')
            doc.add_paragraph(f'Quando você cometer erros \n{serie2_discurso2}')
            doc.add_paragraph(f'Quando os maus forem bem-sucedidos \n{serie2_discurso3}')
            doc.add_paragraph(f'15:40 “Há uma recompensa para o justo \n{discurso5}')
            doc.add_paragraph(f'16:15 Cântico 140 e oração')
            doc.add_heading('Preste atenção às respostas para as seguintes perguntas:', 0)
            doc.add_paragraph(f'1. Como nós ‘esperamos ansiosamente por Jeová’? (Salmo 130:5, 6) \n{recap_perg1}')
            doc.add_paragraph(f'2. Como podemos esperar ansiosamente por Jeová quando enfrentamos dificuldades? (Habacuque 2:3, 4; 2 Timóteo 4:2; Lucas 2:36-38 \n{recap_perg2}')
            doc.add_paragraph(f'3. Como podemos usar de forma sábia o nosso tempo enquanto esperamos pelo dia de Jeová? (2 Pedro 3:11-13) \n{recap_perg3}')
            doc.add_paragraph(f'4. Por que podemos esperar por Jeová com confiança quando enfrentamos desafios? (Salmo 62:1, 2, 8, 10; Capítulo 68:6; Capítulo 130:2 a 4) \n{recap_perg4}')
            doc.add_paragraph(f'5. O que devemos fazer para receber a “recompensa para o justo”? (Salmo 58:11) \n{recap_perg5}')
            doc.add_paragraph(f'6. Como podemos sentir alegria no nosso serviço a Jeová? (João 5:17) \n{recap_perg6}')
        # Salva o documento Word
            doc_file = (nome_arq)
            doc.save(f'Docs\{nome_arq}')

            st.success(f"Documento Word gerado com sucesso, {doc_file}")


def cong():
    
    if tela3:
        titulo = st.title(":blue[Programa da Assembleia de Circuito com o Superintendente de Circuito 2023 a 2024]")
        tema = st.header(":red[Espere Ansiosamente por Jeová!]\n Salmo 130:6")
        manha = st.header(':blue[Manhã]')

        Música = st.text("9:30 Música")
        cantico1 = st.text("9:40 Cântico 88 e oração")
        discurso = st.text_area("9:50 ‘Esperar ansiosamente por Jeová’ — Como?")
        serie_discurso1 = st.text_area("10:05 Série de discursos: \nImite aqueles que esperaram ansiosamente - Habacuque")
        serie_discurso2 = st.text_area("Imite aqueles que esperaram ansiosamente - João")
        serie_discurso3 = st.text_area("Imite aqueles que esperaram ansiosamente - Ana")
        cantico2 = st.text_input("11:05 Cântico 142 e anúncios")
        discurso4 = st.text_area("11:15 O que você está esperando?")
        discurso_batismo = st.text_area("11:30 Dedicação e batismo")
        cantico3 = st.text("12:00 Cântico 28")

        tarde = st.header(':blue[Tarde]')

        musica2 = st.text("13:10 Música")
        cantico4 = st.text("13:20 Cântico 54 e oração")
        discurso_publico = st.text_area("13:30 Discurso para o público: Será que a paciência ainda tem valor?")
        sentinela = st.text_area("14:00 Resumo de A Sentinela")
        cantico5 = st.text_input("14:30 Cântico 143 e anúncios")
        serie2_discurso1 = st.text_area("14:40 Série de discursos: Espere por Jeová \n Quando se sentir sozinho")
        serie2_discurso2 = st.text_area("Quando você cometer erros")
        serie2_discurso3 = st.text_area("Quando os maus forem bem-sucedidos")
        discurso5 = st.text_area("15:40 “Há uma recompensa para o justo")
        cantico6 = st.text("16:15 Cântico 140 e oração")

        st.title(':red[Preste atenção às respostas para as seguintes perguntas:]')

        recap_perg1 = st.text_area("1. Como nós ‘esperamos ansiosamente por Jeová’? (Salmo 130:5, 6)")
        recap_perg2 = st.text_area('2. Como podemos esperar ansiosamente por Jeová quando enfrentamos dificuldades? (Habacuque 2:3, 4; 2 Timóteo 4:2; Lucas 2:36-38')
        recap_perg3 = st.text_area('3. Como podemos usar de forma sábia o nosso tempo enquanto esperamos pelo dia de Jeová? (2 Pedro 3:11-13)')
        recap_perg4 = st.text_area('4. Por que podemos esperar por Jeová com confiança quando enfrentamos desafios? (Salmo 62:1, 2, 8, 10; Capítulo 68:6; Capítulo 130:2 a 4)')
        recap_perg5 = st.text_area('5. O que devemos fazer para receber a “recompensa para o justo”? (Salmo 58:11)')
        nome_arq = st.text_input('Nome do arquivo - no final do nome colocar (.doc)')

        if st.button('Salvar'):
            doc = Document()
            
            doc.add_heading('Programa da Assembleia de Circuito com o Superintendente de Circuito 2023 a 2024', 0)
            doc.add_heading('Espere Ansiosamente por Jeová! - Salmo 130:6', 1)
            doc.add_heading('Manhã', 0)
            doc.add_paragraph('9:30 Música')
            doc.add_paragraph(f'9:50 ‘Esperar ansiosamente por Jeová’ — Como? \n{discurso}')
            doc.add_paragraph(f'10:05 Série de discursos: \nImite aqueles que esperaram ansiosamente - Habacuque \n{serie_discurso1}')
            doc.add_paragraph(f'Imite aqueles que esperaram ansiosamente - João \n{serie_discurso2}')
            doc.add_paragraph(f'Imite aqueles que esperaram ansiosamente - Ana \n{serie_discurso3}')
            doc.add_paragraph(f'11:05 Cântico 142 e anúncios \n{cantico2}')
            doc.add_paragraph(f'11:15 O que você está esperando? \n{discurso4}')
            doc.add_paragraph(f'11:30 Dedicação e batismo \n{discurso_batismo}')
            doc.add_paragraph(f'12:00 Cântico 28')
            
            doc.add_heading('Tarde', 0)
            
            doc.add_paragraph(f'13:10 Música')
            doc.add_paragraph(f'13:20 Cântico 54 e oração')
            doc.add_paragraph(f'13:30 Discurso para o público: Será que a paciência ainda tem valor? \n{discurso_publico}')
            doc.add_paragraph(f'14:00 Resumo de A Sentinela \n{sentinela}')
            doc.add_paragraph(f'14:30 Cântico 143 e anúncios \n{cantico5}')
            doc.add_paragraph(f'14:40 Série de discursos: Espere por Jeová: \n Quando se sentir sozinho \n{serie2_discurso1}')
            doc.add_paragraph(f'Quando você cometer erros \n{serie2_discurso2}')
            doc.add_paragraph(f'Quando os maus forem bem-sucedidos \n{serie2_discurso3}')
            doc.add_paragraph(f'15:40 “Há uma recompensa para o justo \n{discurso5}')
            doc.add_paragraph(f'16:15 Cântico 140 e oração')
            doc.add_heading('Preste atenção às respostas para as seguintes perguntas:', 0)
            doc.add_paragraph(f'1. Como nós ‘esperamos ansiosamente por Jeová’? (Salmo 130:5, 6) \n{recap_perg1}')
            doc.add_paragraph(f'2. Como podemos esperar ansiosamente por Jeová quando enfrentamos dificuldades? (Habacuque 2:3, 4; 2 Timóteo 4:2; Lucas 2:36-38 \n{recap_perg2}')
            doc.add_paragraph(f'3. Como podemos usar de forma sábia o nosso tempo enquanto esperamos pelo dia de Jeová? (2 Pedro 3:11-13) \n{recap_perg3}')
            doc.add_paragraph(f'4. Por que podemos esperar por Jeová com confiança quando enfrentamos desafios? (Salmo 62:1, 2, 8, 10; Capítulo 68:6; Capítulo 130:2 a 4) \n{recap_perg4}')
            doc.add_paragraph(f'5. O que devemos fazer para receber a “recompensa para o justo”? (Salmo 58:11) \n{recap_perg5}')
    
        # Salva o documento Word
            doc_file = (nome_arq)
            doc.save(f'Docs\{nome_arq}')

            st.success(f"Documento Word gerado com sucesso, {doc_file}")

tela1 = st.sidebar.checkbox('Assembléia Superintendente de Circuito')
if tela1:
    asc()
tela2 = st.sidebar.checkbox('Assembléia Representante de Betel')
if tela2:
    arb()
tela3 = st.sidebar.checkbox('Congresso')
# if tela3:
#     cong()
