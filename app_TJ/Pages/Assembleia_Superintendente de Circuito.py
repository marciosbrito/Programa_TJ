import streamlit as st
import docx
from docx import Document

titulo = st.title(":blue[Programa da Assembleia de Circuito com o Superintendente de Circuito 2023 a 2024]")
tema = st.header(":red[Espere Ansiosamente por Jeová!]\n Salmo 130:6")
manha = st.header(':blue[Manhã]')

Música = st.text("9:30 Música")
cantico1 = st.text("9:40 Cântico 88 e oração")
discurso = st.text_area("9:50 ‘Esperar ansiosamente por Jeová’ — Como?")
serie_discurso1 = st.text_area("10:05 Série de discursos: \nImite aqueles que esperaram ansiosamente - Habacuque")
serie_discurso2 = st.text_area("Imite aqueles que esperaram ansiosamente - João")
serie_discurso3 = st.text_area("Imite aqueles que esperaram ansiosamente - Ana")
cantico2 = st.text_input("11:05 Cântico 142 e anúncios")
discurso4 = st.text_area("11:15 O que você está esperando?")
discurso_batismo = st.text_area("11:30 Dedicação e batismo")
cantico3 = st.text("12:00 Cântico 28")

tarde = st.header(':blue[Tarde]')

musica2 = st.text("13:10 Música")
cantico4 = st.text("13:20 Cântico 54 e oração")
discurso_publico = st.text_area("13:30 Discurso para o público: Será que a paciência ainda tem valor?")
sentinela = st.text_area("14:00 Resumo de A Sentinela")
cantico5 = st.text_input("14:30 Cântico 143 e anúncios")
serie2_discurso1 = st.text_area("14:40 Série de discursos: Espere por Jeová \n Quando se sentir sozinho")
serie2_discurso2 = st.text_area("Quando você cometer erros")
serie2_discurso3 = st.text_area("Quando os maus forem bem-sucedidos")
discurso5 = st.text_area("15:40 “Há uma recompensa para o justo")
cantico6 = st.text("16:15 Cântico 140 e oração")

st.title(':red[Preste atenção às respostas para as seguintes perguntas:]')

recap_perg1 = st.text_area("1. Como nós ‘esperamos ansiosamente por Jeová’? (Salmo 130:5, 6)")
recap_perg2 = st.text_area('2. Como podemos esperar ansiosamente por Jeová quando enfrentamos dificuldades? (Habacuque 2:3, 4; 2 Timóteo 4:2; Lucas 2:36-38')
recap_perg3 = st.text_area('3. Como podemos usar de forma sábia o nosso tempo enquanto esperamos pelo dia de Jeová? (2 Pedro 3:11-13)')
recap_perg4 = st.text_area('4. Por que podemos esperar por Jeová com confiança quando enfrentamos desafios? (Salmo 62:1, 2, 8, 10; Capítulo 68:6; Capítulo 130:2 a 4)')
recap_perg5 = st.text_area('5. O que devemos fazer para receber a “recompensa para o justo”? (Salmo 58:11)')
nome_arq = st.text_input('Nome do arquivo - no final do nome colocar (.doc)')

if st.button('Salvar'):
    doc = Document()
    
    doc.add_heading('Programa da Assembleia de Circuito com o Superintendente de Circuito 2023 a 2024', 0)
    doc.add_heading('Espere Ansiosamente por Jeová! - Salmo 130:6', 1)
    doc.add_heading('Manhã', 0)
    doc.add_paragraph('9:30 Música')
    doc.add_paragraph(f'9:50 ‘Esperar ansiosamente por Jeová’ — Como? \n{discurso}')
    doc.add_paragraph(f'10:05 Série de discursos: \nImite aqueles que esperaram ansiosamente - Habacuque \n{serie_discurso1}')
    doc.add_paragraph(f'Imite aqueles que esperaram ansiosamente - João \n{serie_discurso2}')
    doc.add_paragraph(f'Imite aqueles que esperaram ansiosamente - Ana \n{serie_discurso3}')
    doc.add_paragraph(f'11:05 Cântico 142 e anúncios \n{cantico2}')
    doc.add_paragraph(f'11:15 O que você está esperando? \n{discurso4}')
    doc.add_paragraph(f'11:30 Dedicação e batismo \n{discurso_batismo}')
    doc.add_paragraph(f'12:00 Cântico 28')
    
    doc.add_heading('Tarde', 0)
    
    doc.add_paragraph(f'13:10 Música')
    doc.add_paragraph(f'13:20 Cântico 54 e oração')
    doc.add_paragraph(f'13:30 Discurso para o público: Será que a paciência ainda tem valor? \n{discurso_publico}')
    doc.add_paragraph(f'14:00 Resumo de A Sentinela \n{sentinela}')
    doc.add_paragraph(f'14:30 Cântico 143 e anúncios \n{cantico5}')
    doc.add_paragraph(f'14:40 Série de discursos: Espere por Jeová: \n Quando se sentir sozinho \n{serie2_discurso1}')
    doc.add_paragraph(f'Quando você cometer erros \n{serie2_discurso2}')
    doc.add_paragraph(f'Quando os maus forem bem-sucedidos \n{serie2_discurso3}')
    doc.add_paragraph(f'15:40 “Há uma recompensa para o justo \n{discurso5}')
    doc.add_paragraph(f'16:15 Cântico 140 e oração')
    doc.add_heading('Preste atenção às respostas para as seguintes perguntas:', 0)
    doc.add_paragraph(f'1. Como nós ‘esperamos ansiosamente por Jeová’? (Salmo 130:5, 6) \n{recap_perg1}')
    doc.add_paragraph(f'2. Como podemos esperar ansiosamente por Jeová quando enfrentamos dificuldades? (Habacuque 2:3, 4; 2 Timóteo 4:2; Lucas 2:36-38 \n{recap_perg2}')
    doc.add_paragraph(f'3. Como podemos usar de forma sábia o nosso tempo enquanto esperamos pelo dia de Jeová? (2 Pedro 3:11-13) \n{recap_perg3}')
    doc.add_paragraph(f'4. Por que podemos esperar por Jeová com confiança quando enfrentamos desafios? (Salmo 62:1, 2, 8, 10; Capítulo 68:6; Capítulo 130:2 a 4) \n{recap_perg4}')
    doc.add_paragraph(f'5. O que devemos fazer para receber a “recompensa para o justo”? (Salmo 58:11) \n{recap_perg5}')
      
    
    

# Salva o documento Word
    doc_file = (nome_arq)
    doc.save(f'{nome_arq}')

    st.success(f"Documento Word gerado com sucesso, {doc_file}")
        